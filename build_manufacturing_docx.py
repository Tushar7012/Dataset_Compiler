from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output" / "docx"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "manufacturing_service_test_fixture_50_pages.docx"


TOPICS = [
    ("Fixture purpose and scope", "Sets the fictional manufacturing service context and defines safe boundaries for application testing.", "Use synthetic assets, work orders, suppliers, and operators only.", "Keep fixture ID MFG-FIXTURE-50 on imported records.", "Test page counting, section detection, tables, and metadata extraction.", "Fixture owner", "Northstar Manufacturing Lab"),
    ("Plant service model", "Describes how plant operations, maintenance, quality, logistics, and support teams coordinate service work.", "Plant operations owns production continuity.", "Maintenance owns equipment health and work execution.", "Quality owns release criteria and nonconformance disposition.", "Service model", "Integrated plant support"),
    ("Plant and line hierarchy", "Defines the hierarchy used to locate equipment, assign responsibility, and aggregate downtime.", "Plant P-04 contains Lines L-11 through L-16.", "Each line contains cells, stations, assets, and instrument points.", "Hierarchy changes require an effective date and owner.", "Example path", "P-04 / L-12 / Cell-C / Press-07"),
    ("Asset registry standards", "Every maintainable asset has a stable identity, location, criticality, lifecycle state, and service owner.", "Use asset tag AST-P04-L12-007 for parser and search tests.", "Separate manufacturer serial from internal asset tag.", "Retire assets only after history and successor links are preserved.", "Required field", "Asset tag, status, location, criticality"),
    ("Asset criticality scoring", "Criticality ranks equipment by safety, quality, production, environmental, and recovery impact.", "Score safety and regulatory impact before production throughput.", "Review criticality after process or product changes.", "Use the score to set inspection and spare-part priority.", "Criticality", "A1 - production and safety critical"),
    ("Equipment taxonomy", "A controlled taxonomy groups equipment into useful service families without hiding local details.", "Use Press, Conveyor, Robot, Furnace, Pump, and Inspection families.", "Add model and function as separate attributes.", "Do not encode changing location inside the asset ID.", "Family", "Robotic material handling"),
    ("Preventive maintenance program", "Preventive maintenance converts manufacturer guidance and plant history into scheduled tasks.", "Define interval, trigger, skill, tools, and acceptance criteria.", "Record actual completion time separately from planned time.", "Review overdue tasks by criticality and risk.", "Interval", "Every 1,000 operating hours"),
    ("Predictive maintenance signals", "Condition signals help identify degradation before a functional failure occurs.", "Trend vibration, temperature, current, pressure, and cycle time.", "Store signal source and sampling interval.", "Validate alerts against technician findings before tuning thresholds.", "Signal", "Vibration RMS above 7.5 mm/s"),
    ("Corrective maintenance intake", "Corrective work begins from an observed symptom, failure, inspection finding, or production request.", "Capture asset, symptom, time, impact, and safe state.", "Avoid diagnosing root cause in the initial symptom field.", "Link repeat events to a common problem record.", "Work order", "WO-260817-0412"),
    ("Work order lifecycle", "Work orders move through controlled states so planning, execution, review, and closure remain auditable.", "Use New, Approved, Planned, In Progress, Awaiting Parts, Complete, and Closed.", "Require an owner before execution starts.", "Close only after labor, parts, evidence, and failure code are present.", "Lifecycle", "Approved to Planned"),
    ("Work order priority", "Priority combines safety, quality, production loss, customer commitment, and resource availability.", "P0 indicates immediate safety or environmental risk.", "P1 indicates severe production or quality interruption.", "P2 and P3 cover planned and low-impact service work.", "Priority", "P1 - production interruption"),
    ("Maintenance planning board", "The planning board balances backlog, skills, access windows, tools, and production constraints.", "Separate ready backlog from blocked backlog.", "Show planned start and latest safe completion dates.", "Review schedule conflicts with operations before release.", "Planning horizon", "14 calendar days"),
    ("Job plan structure", "A job plan makes work repeatable by defining sequence, hazards, tools, parts, and verification steps.", "Use action verbs and observable completion criteria.", "Place isolation and restoration steps at explicit boundaries.", "Version job plans when equipment or method changes.", "Job plan", "JP-PRESS-014 version 2.1"),
    ("Permit and access control", "Some service work requires a permit, area owner approval, or controlled access window.", "Identify permit type before technician dispatch.", "Store permit ID and approval timestamp on the work order.", "Do not treat a verbal request as permit approval.", "Permit", "Hot work permit HW-0427"),
    ("Lockout tagout workflow", "Energy isolation protects people during service work and must be verified before contact with hazardous energy.", "Identify all energy sources and isolation points.", "Record verifier and zero-energy test result.", "Restore guards and remove tags through the approved release path.", "Safety record", "LOTO-P04-260817-09"),
    ("Confined space service", "Confined space work requires entry criteria, atmospheric testing, attendant coverage, and rescue readiness.", "Confirm space classification and current permit.", "Record gas readings with time and instrument ID.", "Stop entry when conditions leave approved limits.", "Atmosphere", "O2 20.8%, LEL 0%, H2S 0 ppm"),
    ("Electrical service controls", "Electrical tasks use competency, voltage verification, arc-flash boundaries, and approved test equipment.", "Match technician authorization to task voltage.", "Verify test instrument status before use.", "Document restoration and functional check.", "Control", "Qualified person and voltage test"),
    ("Mechanical repair standards", "Mechanical repairs preserve fit, alignment, torque, lubrication, and guarding requirements.", "Use approved torque and clearance values.", "Record replaced component and serial when traceability matters.", "Run a controlled test cycle before handback.", "Acceptance", "Guarding restored and test cycle passed"),
    ("Robot and automation service", "Robot service requires program version, safe mode, teach pendant control, and recovery verification.", "Capture controller and program revision before change.", "Use approved backup before editing motion logic.", "Validate home position, interlocks, and dry cycle.", "Controller", "RB-12 / Program 6.4.2"),
    ("Calibration management", "Calibration records establish whether instruments were fit for use during a measurement period.", "Track instrument ID, due date, standard, and result.", "Quarantine failed instruments and assess affected work.", "Link calibration certificate to the service record.", "Instrument", "CAL-TEMP-118"),
    ("Quality inspection planning", "Inspection plans define what to measure, when to measure, method, sample, and reaction plan.", "Use characteristic ID rather than free-text names only.", "Separate measurement value from pass or fail decision.", "Escalate out-of-control patterns before release.", "Characteristic", "CTQ-07 hole diameter"),
    ("Nonconformance intake", "Nonconformance records capture defect, containment, affected material, and disposition ownership.", "Use lot, batch, serial, and operation references.", "Contain first when product identity is uncertain.", "Keep disposition approval separate from detection.", "NCR", "NCR-260817-033"),
    ("Root cause analysis", "Root cause work moves from symptom to causal evidence and verifies that proposed controls address recurrence.", "Use 5 Why, fishbone, or fault tree based on complexity.", "Distinguish contributing factor from verified cause.", "Define evidence and due date for corrective action.", "Method", "5 Why with process confirmation"),
    ("Corrective and preventive action", "CAPA tracks actions, owners, due dates, evidence, and effectiveness checks.", "Write actions as observable changes, not intentions.", "Require risk review when the action changes process controls.", "Close effectiveness only after representative monitoring.", "CAPA", "CAPA-260817-014"),
    ("Production changeover support", "Changeover service coordinates tooling, recipe, cleaning, inspection, and first-off approval.", "Confirm previous product clearance.", "Record setup parameters and first-off result.", "Escalate if setup time or quality limit is exceeded.", "Changeover", "Line L-12 product family B to C"),
    ("Tooling and fixture service", "Tooling records connect fixture identity, condition, calibration, storage, and repair history.", "Track tool life and inspection interval.", "Quarantine damaged or overdue fixtures.", "Confirm correct revision at point of use.", "Tool tag", "FIX-L12-044"),
    ("Spare parts catalog", "The parts catalog supports identification, alternates, stock policy, and work order consumption.", "Use manufacturer part number and internal item number.", "Mark approved alternates with effective dates.", "Prevent duplicate items created from naming variation.", "Part", "PN-VALVE-4407"),
    ("Inventory and min-max policy", "Inventory controls protect service readiness without creating unnecessary capital or obsolete stock.", "Set minimum and maximum by lead time and criticality.", "Review demand from planned and corrective work.", "Record stockout reason when work is delayed.", "Stock rule", "Min 2, Max 6, critical A1"),
    ("Parts reservation and issue", "Reserved parts move from available stock to a named work order with clear issue and return states.", "Reserve only after part identity is verified.", "Record quantity, lot, and serial where applicable.", "Return unused parts with condition and reason.", "Issue", "ISS-260817-0098"),
    ("Supplier service management", "Supplier records cover scope, contact path, service level, certifications, and performance evidence.", "Store contract reference separately from supplier display name.", "Validate certifications before assigning regulated work.", "Review performance using completed service evidence.", "Supplier", "SUP-118 / Apex Motion Services"),
    ("External technician onboarding", "External technicians receive access, safety orientation, site rules, and task boundaries before work begins.", "Verify identity and qualification expiry.", "Issue least-privilege system access.", "Revoke access when assignment ends.", "Access", "Badge B-4431 active through 2026-09-30"),
    ("Service level and response targets", "Service targets define response, arrival, restore, and report timelines by priority and contract.", "Keep target clock rules explicit.", "Pause only for documented customer or site dependency.", "Measure breach reasons, not only breach counts.", "P1 target", "Arrival within 60 minutes"),
    ("Downtime and availability", "Downtime records distinguish planned, unplanned, blocked, and administrative time for reliable availability analysis.", "Capture start, end, reason, asset, and production impact.", "Avoid double-counting overlapping work orders.", "Reconcile event log with line status.", "Availability", "Run time divided by planned production time"),
    ("Overall equipment effectiveness", "OEE combines availability, performance, and quality to show how equipment converts planned time into good output.", "Use a consistent planned production denominator.", "Trace losses to event or quality records.", "Do not compare lines with different calculation rules.", "OEE loss", "Availability loss from unplanned stop"),
    ("Reliability review", "Reliability review converts failure history into prioritized improvement work.", "Trend MTBF, MTTR, repeat failure, and infant mortality.", "Segment by asset family and operating context.", "Validate outliers before changing maintenance strategy.", "Metric", "MTTR 42 minutes for Press-07"),
    ("Maintenance cost reporting", "Cost views combine labor, parts, contractor, downtime, and planned improvement spend.", "Allocate cost to asset and work type.", "Separate estimate, commitment, and actual cost.", "Explain unusual variance with evidence.", "Cost center", "CC-P04-MAINT"),
    ("Environmental service controls", "Environmental controls cover spills, emissions equipment, waste handling, and regulated maintenance evidence.", "Identify environmental impact in work classification.", "Record response action and notification path.", "Retain inspection evidence for required duration.", "Event", "Minor coolant leak contained"),
    ("Waste and hazardous material handling", "Waste handling preserves labeling, container state, pickup record, and approved disposal route.", "Classify material before movement.", "Keep container and manifest references.", "Escalate damaged containers immediately.", "Waste stream", "Used solvent, regulated"),
    ("Emergency maintenance response", "Emergency service prioritizes life safety, containment, communication, and controlled restoration.", "Stop work when conditions are unsafe.", "Name incident commander and technical lead.", "Document temporary repair and follow-up permanent fix.", "Emergency", "EMG-P04-260817-02"),
    ("Shift handover", "Handover records preserve open risks, isolated assets, blocked work, permits, and next decisions.", "Use structured handover fields.", "Call out changes since previous shift.", "Require receiving shift acknowledgment.", "Handover", "Shift A to Shift B at 18:00"),
    ("Technician competency matrix", "Competency data links people to tasks they may perform independently, with supervision, or not at all.", "Track skill, level, evidence, and expiry.", "Use matrix during planning and dispatch.", "Do not infer qualification from job title.", "Skill", "Electrical isolation level 2"),
    ("Service documentation standard", "Service documents should allow a second technician to understand what changed and how the asset was verified.", "Use stable asset and work order identifiers.", "Attach photos or readings with timestamps.", "Avoid unexplained abbreviations in closeout notes.", "Document", "WO-260817-0412 closeout"),
    ("Data quality controls", "Data quality checks find missing, invalid, duplicate, stale, and conflicting operational records.", "Validate mandatory fields at transition points.", "Use exception queues for records that cannot be auto-fixed.", "Track defect source and correction outcome.", "Rule", "DQ-MFG-017 missing asset location"),
    ("Mobile service execution", "Mobile execution keeps technicians productive while preserving reliable work order, evidence, and safety records at the point of work.", "Cache only the minimum data needed for an approved offline window.", "Queue updates with event IDs and resolve conflicts explicitly.", "Confirm synchronization status before closing work.", "Device", "MOB-P04-TECH-22"),
    ("Integration event model", "Integration events synchronize work orders, assets, inventory, and quality findings across plant systems.", "Use event ID and retry key for idempotency.", "Preserve source timestamp and processing timestamp.", "Send rejected events to a reviewable dead-letter path.", "Event", "work_order.status.changed"),
    ("Search and indexing anchors", "Stable anchors support full-text, metadata, and semantic retrieval tests across this fixture.", "Search for MFG-FIXTURE-50 and AST-P04-L12-007.", "Search for WO-260817-0412 and NCR-260817-033.", "Search for Lockout tagout and preventive maintenance.", "Document ID", "MFG-FIXTURE-50"),
    ("Appendix: sample records", "Sample records create repeatable cross-reference tests between assets, work orders, parts, suppliers, and quality cases.", "AST-P04-L12-007 is linked to WO-260817-0412.", "WO-260817-0412 consumes PN-VALVE-4407.", "NCR-260817-033 references Line L-12.", "Record set", "MFG-SAMPLE-014"),
    ("Appendix: status and reason codes", "Reference codes keep workflow fixtures consistent across imported data and analytics.", "NEW means newly created.", "BLK means blocked by dependency.", "CLS means closed after verification.", "Code", "WO-NEW, WO-BLK, WO-CLS"),
    ("Appendix: acceptance checklist", "Final checklist defines what an application should be able to validate after ingesting this document.", "Confirm document has 50 rendered pages.", "Confirm headings and tables remain searchable.", "Confirm fictional identifiers remain intact.", "Acceptance", "50 pages, no clipping, stable anchors"),
    ("Appendix: fixture metadata", "Captures stable metadata for automated regression tests and comparison across application versions.", "Title is Manufacturing Service Test Fixture.", "Revision is 1.0 and effective date is 2026-08-17.", "All content is synthetic and non-production.", "Pages", "50"),
]


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[index]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = trPr.first_child_found_in("w:tblHeader")
    if tblHeader is None:
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
    tblHeader.set(qn("w:val"), "true")


def set_run_font(run, name="Calibri", size=11, color="243447", bold=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("243447")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [("Heading 1", 16, "2E74B5", 18, 10), ("Heading 2", 13, "2E74B5", 14, 7), ("Heading 3", 12, "1F4D78", 10, 5)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = doc.styles["List Bullet"]
    list_style.font.name = "Calibri"
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = Inches(-0.188)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("NORTHSTAR MANUFACTURING LAB | TEST FIXTURE"), size=8, color="51606F", bold=True)
    pPr = header._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D7E0EA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("Manufacturing Service Manual | Page "), size=8, color="51606F")
    add_page_number(footer)

    for index, (title, overview, b1, b2, b3, label, value) in enumerate(TOPICS, start=1):
        kicker = doc.add_paragraph()
        kicker.paragraph_format.space_after = Pt(8)
        set_run_font(kicker.add_run(f"MANUFACTURING SERVICE | MODULE {index:02d}"), size=8.5, color="2E74B5", bold=True)

        heading = doc.add_paragraph(style="Heading 1")
        set_run_font(heading.add_run(title), size=16, color="2E74B5", bold=True)

        body = doc.add_paragraph()
        body.paragraph_format.space_after = Pt(8)
        body.add_run(overview)

        label_p = doc.add_paragraph(style="Heading 2")
        label_p.paragraph_format.space_before = Pt(8)
        label_p.paragraph_format.space_after = Pt(5)
        set_run_font(label_p.add_run("Execution controls"), size=13, color="2E74B5", bold=True)
        for text in (b1, b2, b3):
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(text)

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        set_table_geometry(table, [2700, 6660])
        mark_header_row(table.rows[0])
        values = [("Field", "Example"), (label, value), ("Fixture ID", "MFG-FIXTURE-50"), ("Module", f"{index:02d} of 50")]
        for r, row_values in enumerate(values):
            for c, text in enumerate(row_values):
                cell = table.cell(r, c)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                set_run_font(run, size=9.5 if r else 9.5, color="0B2545" if r == 0 else "51606F", bold=(r == 0))
                if r == 0:
                    shade_cell(cell, "E8EEF5")
        marker = doc.add_paragraph()
        marker.paragraph_format.space_before = Pt(8)
        marker.paragraph_format.space_after = Pt(0)
        set_run_font(marker.add_run(f"Fixture marker: MFG-FIXTURE-50 | Module {index:02d} | Synthetic content for application testing"), size=8.5, color="51606F")
        if index != len(TOPICS):
            doc.add_page_break()

    doc.core_properties.title = "Manufacturing Service Test Fixture"
    doc.core_properties.subject = "Synthetic 50-page manufacturing service document for application testing"
    doc.core_properties.author = "Northstar Manufacturing Lab"
    doc.save(DOCX_PATH)
    print(f"created {DOCX_PATH}")


if __name__ == "__main__":
    build()
